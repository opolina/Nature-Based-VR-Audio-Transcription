#!/usr/bin/env python3
"""
WhisperX Audio Transcription Tool

Transcribes audio files using WhisperX with word-level timestamp alignment.
Produces multiple export formats (TXT, CSV, XLSX, DOCX) for temporal
correlation with concurrent recordings (e.g., screen capture, eye tracking).

Usage:
    python transcribe.py                              # opens GUI
    python transcribe.py recording.wav                # single file
    python transcribe.py --batch audio_dir/ -o results/
"""

import os
import sys
import gc
import csv
import json
import time
import tkinter as tk
import argparse
import subprocess
import threading
from tkinter import ttk, filedialog, messagebox
from types import SimpleNamespace

import torch
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

# WhisperX's bundled models were saved with older PyTorch defaults.
# Recent PyTorch versions require weights_only=True by default, which
# breaks loading these checkpoints. This patch restores the old behavior.
_original_torch_load = torch.load
def _patched_torch_load(*args, **kwargs):
    kwargs.setdefault("weights_only", False)
    return _original_torch_load(*args, **kwargs)
torch.load = _patched_torch_load

import whisperx

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SAMPLE_RATE = 16000
AUDIO_EXTENSIONS = {".wav", ".mp3", ".flac", ".ogg", ".m4a", ".wma", ".aac", ".opus"}
AUDIO_FILETYPES = [
    ("Audio files", " ".join(f"*{ext}" for ext in sorted(AUDIO_EXTENSIONS))),
    ("All files", "*.*"),
]
MODEL_CHOICES = ["large-v3", "large-v2", "medium", "small", "base", "tiny"]
LANGUAGE_CHOICES = [
    "en", "es", "fr", "de", "it", "pt", "nl", "ru", "zh", "ja",
    "ko", "ar", "hi", "tr", "pl", "sv", "da", "fi", "no", "cs",
    "el", "he", "hu", "id", "ms", "ro", "sk", "th", "uk", "vi",
]
ALL_FORMATS = ["txt", "csv_segments", "csv_words", "xlsx", "docx"]
DEFAULT_FORMATS = {"txt", "csv_segments", "xlsx"}

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SETTINGS_FILE = os.path.join(SCRIPT_DIR, ".transcribe_settings.json")

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

_gui_log_widget = None

def log(msg):
    print(msg)
    w = _gui_log_widget
    if w is not None:
        w.configure(state="normal")
        w.insert("end", msg + "\n")
        w.see("end")
        w.configure(state="disabled")

# ---------------------------------------------------------------------------
# Settings persistence (GUI only)
# ---------------------------------------------------------------------------

def load_settings():
    try:
        with open(SETTINGS_FILE, "r") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def save_settings(settings):
    try:
        with open(SETTINGS_FILE, "w") as f:
            json.dump(settings, f, indent=2)
    except OSError:
        pass

# ---------------------------------------------------------------------------
# Device detection
# ---------------------------------------------------------------------------

def detect_device():
    if torch.cuda.is_available():
        name = torch.cuda.get_device_name(0)
        vram = torch.cuda.get_device_properties(0).total_memory / (1024 ** 3)
        log(f"[INFO] GPU: {name} ({vram:.1f} GB VRAM)")
        return "cuda", "float16" if vram >= 6.0 else "int8"
    log("[WARN] No CUDA GPU detected — falling back to CPU (slow).")
    return "cpu", "int8"

def _free_gpu(device):
    gc.collect()
    if device == "cuda":
        torch.cuda.empty_cache()

# ---------------------------------------------------------------------------
# Word-level data helpers
# ---------------------------------------------------------------------------

def _word_row(word):
    """Extract (start, end, duration, confidence) from a word dict, handling missing values."""
    s, e = word.get("start"), word.get("end")
    dur = round(e - s, 3) if s is not None and e is not None else None
    score = word.get("score")
    return (
        round(s, 3) if s is not None else None,
        round(e, 3) if e is not None else None,
        dur,
        round(score, 3) if score is not None else None,
    )

def _format_ts(seconds):
    return f"{int(seconds // 60)}:{seconds % 60:05.2f}"

# ---------------------------------------------------------------------------
# Export generators
# ---------------------------------------------------------------------------

def generate_txt(segments, path):
    with open(path, "w", encoding="utf-8") as f:
        for seg in segments:
            f.write(seg["text"].strip() + "\n")
    log(f"[INFO] TXT written: {path}")


def generate_csv_segments(segments, path):
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["segment", "start", "end", "duration", "text"])
        for i, seg in enumerate(segments, 1):
            s, e = seg["start"], seg["end"]
            w.writerow([i, round(s, 3), round(e, 3), round(e - s, 3), seg["text"].strip()])
    log(f"[INFO] CSV (segments) written: {path}")


def generate_csv_words(segments, path):
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["segment", "word_index", "word", "start", "end", "duration", "confidence"])
        idx = 0
        for seg_i, seg in enumerate(segments, 1):
            for word in seg.get("words", []):
                idx += 1
                s, e, dur, conf = _word_row(word)
                w.writerow([seg_i, idx, word["word"],
                            s if s is not None else "",
                            e if e is not None else "",
                            dur if dur is not None else "",
                            conf if conf is not None else ""])
    log(f"[INFO] CSV (words) written: {path}")


def generate_xlsx(segments, path):
    wb = Workbook()
    hfont = Font(bold=True)
    hfill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

    def _write_header(ws, headers, row=1):
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font, c.fill = hfont, hfill

    # Segments sheet
    ws = wb.active
    ws.title = "Segments"
    ws.cell(row=1, column=1,
            value="Sheets: Segments, Words, Full Transcript (see tabs below).").font = Font(italic=True, color="666666")
    ws.merge_cells("A1:E1")
    _write_header(ws, ["Segment", "Start (s)", "End (s)", "Duration (s)", "Text"], row=2)
    for i, seg in enumerate(segments, 1):
        s, e = seg["start"], seg["end"]
        ws.append([i, round(s, 3), round(e, 3), round(e - s, 3), seg["text"].strip()])
    ws.column_dimensions["E"].width = 80

    # Words sheet
    ws2 = wb.create_sheet("Words")
    _write_header(ws2, ["Segment", "Word #", "Word", "Start (s)", "End (s)", "Duration (s)", "Confidence"])
    idx = 0
    for seg_i, seg in enumerate(segments, 1):
        for word in seg.get("words", []):
            idx += 1
            s, e, dur, conf = _word_row(word)
            ws2.append([seg_i, idx, word["word"], s, e, dur, conf])

    # Full Transcript sheet
    ws3 = wb.create_sheet("Full Transcript")
    ws3.cell(row=1, column=1, value="Full Transcript").font = hfont
    ws3.cell(row=2, column=1,
             value=" ".join(seg["text"].strip() for seg in segments)).alignment = Alignment(wrap_text=True)
    ws3.column_dimensions["A"].width = 120

    wb.save(path)
    log(f"[INFO] XLSX written: {path}")


def generate_docx(segments, path):
    doc = Document()
    doc.add_heading("Audio Transcription", level=1)
    doc.add_heading("Full Transcript", level=2)
    doc.add_paragraph(" ".join(seg["text"].strip() for seg in segments))

    doc.add_heading("Timestamped Segments", level=2)
    table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
    for cell, text in zip(table.rows[0].cells, ["#", "Start", "End", "Text"]):
        cell.text = text
    for i, seg in enumerate(segments, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = _format_ts(seg["start"])
        row[2].text = _format_ts(seg["end"])
        row[3].text = seg["text"].strip()

    doc.save(path)
    log(f"[INFO] DOCX written: {path}")


# Map format names to (generator_function, filename_suffix)
_FORMAT_TABLE = {
    "txt":          (generate_txt,          "{base}.txt"),
    "csv_segments": (generate_csv_segments, "{base}_segments.csv"),
    "csv_words":    (generate_csv_words,    "{base}_words.csv"),
    "xlsx":         (generate_xlsx,         "{base}.xlsx"),
    "docx":         (generate_docx,         "{base}.docx"),
}

def write_outputs(segments, output_dir, base_name, formats):
    os.makedirs(output_dir, exist_ok=True)
    for fmt in formats:
        gen, pattern = _FORMAT_TABLE[fmt]
        gen(segments, os.path.join(output_dir, pattern.format(base=base_name)))

# ---------------------------------------------------------------------------
# Core transcription pipeline
# ---------------------------------------------------------------------------

def transcribe_audio(audio_path, args):
    if not os.path.isfile(audio_path):
        raise FileNotFoundError(f"Audio file not found: {audio_path}")

    device = args.device or detect_device()[0]
    compute_type = args.compute_type or detect_device()[1]
    if not args.device:
        device, compute_type = detect_device()

    log(f"[INFO] Processing: {audio_path}")
    log(f"[INFO] Device={device}  Compute={compute_type}  Model={args.model}")

    audio = whisperx.load_audio(audio_path)
    log(f"[INFO] Audio duration: {len(audio) / SAMPLE_RATE:.1f}s")

    # ASR options
    asr_options = {
        "no_speech_threshold": 0.3,
        "log_prob_threshold": -0.5,
        "compression_ratio_threshold": 1.8,
        "condition_on_previous_text": False,
    }
    if args.initial_prompt:
        asr_options["initial_prompt"] = args.initial_prompt
        log(f"[INFO] Prompt: {args.initial_prompt[:80]}{'...' if len(args.initial_prompt) > 80 else ''}")
    if args.high_accuracy:
        asr_options["beam_size"] = 10
        asr_options["patience"] = 2.0
        log("[INFO] High accuracy mode (beam_size=10, patience=2.0)")

    chunk_size = args.chunk_size
    vad_options = {"vad_onset": 0.75, "vad_offset": 0.5, "chunk_size": chunk_size}

    log(f"[INFO] Loading model '{args.model}'...")
    model = whisperx.load_model(
        args.model, device,
        compute_type=compute_type,
        language=args.language,
        asr_options=asr_options,
        vad_options=vad_options,
    )

    # Transcribe (auto-retry on GPU OOM with batch_size=1)
    log(f"[INFO] Transcribing (batch_size={args.batch_size}, chunk_size={chunk_size}s)...")
    t0 = time.time()
    try:
        result = model.transcribe(audio, batch_size=args.batch_size, chunk_size=chunk_size)
    except torch.cuda.OutOfMemoryError:
        log(f"[WARN] GPU OOM — retrying with batch_size=1...")
        _free_gpu(device)
        result = model.transcribe(audio, batch_size=1, chunk_size=chunk_size)
    log(f"[INFO] Transcription done in {time.time() - t0:.1f}s")

    lang = args.language or result.get("language", "en")
    log(f"[INFO] Language: {lang}")

    del model
    _free_gpu(device)

    # Word-level alignment
    try:
        log(f"[INFO] Loading alignment model for '{lang}'...")
        model_a, metadata = whisperx.load_align_model(language_code=lang, device=device)
        t0 = time.time()
        result = whisperx.align(
            result["segments"], model_a, metadata, audio, device,
            return_char_alignments=False,
        )
        log(f"[INFO] Alignment done in {time.time() - t0:.1f}s")
        del model_a
        _free_gpu(device)
    except ValueError as e:
        log(f"[WARN] No alignment model for '{lang}': {e}")
        log("[WARN] Skipping word-level alignment — segment timestamps only.")

    # Write outputs
    base_name = os.path.splitext(os.path.basename(audio_path))[0]
    write_outputs(result["segments"], args.output_dir, base_name, args.formats)
    log(f"[INFO] Done: {audio_path}")
    return result

# ---------------------------------------------------------------------------
# Batch helpers
# ---------------------------------------------------------------------------

def _find_audio_files(directory):
    return sorted(
        os.path.join(directory, f) for f in os.listdir(directory)
        if os.path.splitext(f)[1].lower() in AUDIO_EXTENSIONS
    )

def _run_file_subprocess(audio_path, args):
    """Run transcription in an isolated subprocess so the OS reclaims all
    GPU memory on exit — avoids CUDA OOM when CTranslate2 doesn't fully
    release VRAM between consecutive runs."""
    cmd = [sys.executable, "-u", os.path.abspath(__file__), audio_path,
           "-o", args.output_dir, "-m", args.model, "-l", args.language,
           "--batch-size", str(args.batch_size), "--chunk-size", str(args.chunk_size),
           "--formats"] + sorted(args.formats)
    if args.initial_prompt:
        cmd.extend(["--initial-prompt", args.initial_prompt])
    if args.high_accuracy:
        cmd.append("--high-accuracy")

    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
    for line in proc.stdout:
        log(line.rstrip("\n"))
    proc.wait()
    return proc.returncode == 0

def run_batch(file_paths, args):
    """Process multiple files, logging progress and failures."""
    n = len(file_paths)
    if n == 0:
        log("[WARN] No audio files to process.")
        return
    log(f"[INFO] Processing {n} audio file(s)")
    failed = []
    for i, path in enumerate(file_paths, 1):
        if n > 1:
            log(f"\n{'=' * 60}")
            log(f"[INFO] File {i}/{n}: {os.path.basename(path)}")
            log(f"{'=' * 60}")
        if not _run_file_subprocess(path, args):
            failed.append(os.path.basename(path))
    if n > 1:
        log(f"\n{'=' * 60}")
        log(f"[INFO] Batch complete: {n - len(failed)}/{n} succeeded")
        if failed:
            log(f"[WARN] Failed: {', '.join(failed)}")
    log("\n[DONE] All files processed.")

# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

def run_gui():
    global _gui_log_widget

    settings = load_settings()

    root = tk.Tk()
    root.title("Audio Transcription Tool")
    root.resizable(True, True)
    root.minsize(680, 620)

    # -- State variables --
    input_mode = tk.StringVar(value="files")
    selected_files = tk.Variable(value=[])
    selected_folder = tk.StringVar(value="")
    output_dir = tk.StringVar(value=settings.get("output_dir", SCRIPT_DIR))
    model_var = tk.StringVar(value=settings.get("model", "large-v3"))
    language_var = tk.StringVar(value=settings.get("language", "en"))
    batch_size_var = tk.IntVar(value=settings.get("batch_size", 16))
    chunk_size_var = tk.IntVar(value=settings.get("chunk_size", 5))
    high_accuracy_var = tk.BooleanVar(value=settings.get("high_accuracy", False))
    initial_prompt_var = tk.StringVar(value=settings.get("initial_prompt", ""))
    is_running = False

    format_vars = {}
    saved_fmts = settings.get("formats", list(DEFAULT_FORMATS))
    for fmt in ALL_FORMATS:
        format_vars[fmt] = tk.BooleanVar(value=(fmt in saved_fmts))

    # -- Input section --
    input_frame = ttk.LabelFrame(root, text="Input", padding=10)
    input_frame.pack(fill="x", padx=10, pady=(10, 5))

    mode_frame = ttk.Frame(input_frame)
    mode_frame.pack(fill="x")
    ttk.Radiobutton(mode_frame, text="Select file(s)", variable=input_mode, value="files").pack(side="left")
    ttk.Radiobutton(mode_frame, text="Select folder (batch)", variable=input_mode, value="folder").pack(side="left", padx=(15, 0))

    path_frame = ttk.Frame(input_frame)
    path_frame.pack(fill="x", pady=(5, 0))
    path_label = ttk.Label(path_frame, text="No files selected", foreground="gray")
    path_label.pack(side="left", fill="x", expand=True)

    def browse_input():
        if input_mode.get() == "files":
            paths = filedialog.askopenfilenames(title="Select audio file(s)", filetypes=AUDIO_FILETYPES)
            if paths:
                selected_files.set(list(paths))
                selected_folder.set("")
                names = ", ".join(os.path.basename(p) for p in paths[:3])
                suffix = f" (+{len(paths) - 3} more)" if len(paths) > 3 else ""
                path_label.config(text=f"{names}{suffix}", foreground="black")
        else:
            folder = filedialog.askdirectory(title="Select folder with audio files")
            if folder:
                selected_folder.set(folder)
                selected_files.set([])
                path_label.config(text=folder, foreground="black")

    ttk.Button(path_frame, text="Browse...", command=browse_input).pack(side="right")

    # -- Output directory --
    output_frame = ttk.LabelFrame(root, text="Output Directory", padding=10)
    output_frame.pack(fill="x", padx=10, pady=5)
    ttk.Entry(output_frame, textvariable=output_dir).pack(side="left", fill="x", expand=True)

    def browse_output():
        folder = filedialog.askdirectory(title="Select output directory")
        if folder:
            output_dir.set(folder)

    ttk.Button(output_frame, text="Browse...", command=browse_output).pack(side="right", padx=(5, 0))

    # -- Settings --
    settings_frame = ttk.LabelFrame(root, text="Settings", padding=10)
    settings_frame.pack(fill="x", padx=10, pady=5)

    row1 = ttk.Frame(settings_frame)
    row1.pack(fill="x")
    ttk.Label(row1, text="Model:").pack(side="left")
    ttk.Combobox(row1, textvariable=model_var, values=MODEL_CHOICES, state="readonly", width=12).pack(side="left", padx=(5, 20))
    ttk.Label(row1, text="Language:").pack(side="left")
    ttk.Combobox(row1, textvariable=language_var, values=LANGUAGE_CHOICES, state="readonly", width=6).pack(side="left", padx=(5, 20))
    ttk.Label(row1, text="Max segment (s):").pack(side="left")
    ttk.Spinbox(row1, from_=1, to=30, textvariable=chunk_size_var, width=4).pack(side="left", padx=(5, 0))

    row2 = ttk.Frame(settings_frame)
    row2.pack(fill="x", pady=(5, 0))
    ttk.Checkbutton(row2, text="High accuracy", variable=high_accuracy_var).pack(side="left")
    ttk.Label(row2, text="(slower)", foreground="gray").pack(side="left", padx=(5, 0))

    row3 = ttk.Frame(settings_frame)
    row3.pack(fill="x", pady=(5, 0))
    ttk.Label(row3, text="Prompt (optional):").pack(side="left")
    ttk.Entry(row3, textvariable=initial_prompt_var).pack(side="left", fill="x", expand=True, padx=(5, 0))

    # -- Export formats --
    export_frame = ttk.LabelFrame(root, text="Export Formats", padding=10)
    export_frame.pack(fill="x", padx=10, pady=5)

    seg_frame = ttk.LabelFrame(export_frame, text="Segment-level", padding=(8, 4))
    seg_frame.pack(fill="x")
    for fmt, label in [("txt", "Plain text (.txt)"), ("csv_segments", "CSV (.csv)"), ("docx", "Word (.docx)")]:
        ttk.Checkbutton(seg_frame, text=label, variable=format_vars[fmt]).pack(side="left", padx=(0, 12))

    word_frame = ttk.LabelFrame(export_frame, text="Word-level (includes segments)", padding=(8, 4))
    word_frame.pack(fill="x", pady=(5, 0))
    for fmt, label in [("csv_words", "CSV (.csv)"), ("xlsx", "Excel workbook (.xlsx)")]:
        ttk.Checkbutton(word_frame, text=label, variable=format_vars[fmt]).pack(side="left", padx=(0, 12))

    ttk.Label(root,
              text="Note: Automated transcription is not 100% accurate. Please review output for errors.",
              foreground="gray", font=("Segoe UI", 8, "italic")).pack(padx=10, pady=(2, 0), anchor="w")

    # -- Transcribe button & progress --
    btn_frame = ttk.Frame(root)
    btn_frame.pack(fill="x", padx=10, pady=5)
    progress = ttk.Progressbar(btn_frame, mode="indeterminate")
    progress.pack(side="left", fill="x", expand=True, padx=(0, 10))

    def start_transcription():
        nonlocal is_running
        if is_running:
            return

        files = list(selected_files.get())
        folder = selected_folder.get()
        if input_mode.get() == "files" and not files:
            return messagebox.showwarning("No input", "Please select one or more audio files.")
        if input_mode.get() == "folder" and not folder:
            return messagebox.showwarning("No input", "Please select a folder.")
        if not output_dir.get().strip():
            return messagebox.showwarning("No output", "Please specify an output directory.")

        sel_fmts = {fmt for fmt, var in format_vars.items() if var.get()}
        if not sel_fmts:
            return messagebox.showwarning("No formats", "Please select at least one export format.")

        save_settings({
            "output_dir": output_dir.get().strip(), "model": model_var.get(),
            "language": language_var.get().strip(), "batch_size": batch_size_var.get(),
            "chunk_size": chunk_size_var.get(), "formats": list(sel_fmts),
            "high_accuracy": high_accuracy_var.get(),
            "initial_prompt": initial_prompt_var.get().strip(),
        })

        is_running = True
        transcribe_btn.config(state="disabled")
        progress.start(15)
        log_text.configure(state="normal")
        log_text.delete("1.0", "end")
        log_text.configure(state="disabled")

        args = SimpleNamespace(
            model=model_var.get(), language=language_var.get().strip() or "en",
            batch_size=batch_size_var.get(), chunk_size=chunk_size_var.get(),
            output_dir=output_dir.get().strip(), device=None, compute_type=None,
            formats=sel_fmts, high_accuracy=high_accuracy_var.get(),
            initial_prompt=initial_prompt_var.get().strip() or None,
        )

        def run():
            nonlocal is_running
            try:
                if input_mode.get() == "folder":
                    all_files = _find_audio_files(folder)
                else:
                    all_files = list(files)
                run_batch(all_files, args)
            except Exception as e:
                log(f"\n[ERROR] {e}")
            finally:
                is_running = False
                root.after(0, lambda: transcribe_btn.config(state="normal"))
                root.after(0, progress.stop)

        threading.Thread(target=run, daemon=True).start()

    transcribe_btn = ttk.Button(btn_frame, text="Transcribe", command=start_transcription)
    transcribe_btn.pack(side="right")

    # -- Log panel --
    log_frame = ttk.LabelFrame(root, text="Log", padding=5)
    log_frame.pack(fill="both", expand=True, padx=10, pady=(5, 10))

    log_text = tk.Text(log_frame, height=10, wrap="word", state="disabled", font=("Consolas", 9))
    scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=log_text.yview)
    log_text.configure(yscrollcommand=scrollbar.set)
    scrollbar.pack(side="right", fill="y")
    log_text.pack(fill="both", expand=True)

    _gui_log_widget = log_text

    def show_device_info():
        if torch.cuda.is_available():
            name = torch.cuda.get_device_name(0)
            vram = torch.cuda.get_device_properties(0).total_memory / (1024 ** 3)
            log(f"[INFO] Ready. GPU: {name} ({vram:.1f} GB)")
        else:
            log("[INFO] Ready. No GPU detected — will use CPU (slow).")

    root.after(100, show_device_info)
    root.mainloop()

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main_cli():
    parser = argparse.ArgumentParser(
        description="Transcribe audio files using WhisperX with word-level timestamps.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="Examples:\n"
               "  python transcribe.py recording.wav\n"
               "  python transcribe.py recording.wav -o results/\n"
               "  python transcribe.py --batch audio_files/ -o results/\n"
               "  python transcribe.py recording.wav --formats txt csv_segments xlsx\n",
    )

    input_group = parser.add_mutually_exclusive_group(required=True)
    input_group.add_argument("audio_file", nargs="?", help="Path to a single audio file")
    input_group.add_argument("--batch", metavar="DIR", help="Directory for batch processing")

    parser.add_argument("-o", "--output-dir", default=SCRIPT_DIR, help="Output directory (default: script dir)")
    parser.add_argument("-m", "--model", default="large-v3", choices=MODEL_CHOICES, help="Whisper model (default: large-v3)")
    parser.add_argument("-l", "--language", default="en", help="Language code (default: en)")
    parser.add_argument("--batch-size", type=int, default=16, help="Batch size (reduce if GPU OOM, default: 16)")
    parser.add_argument("--chunk-size", type=int, default=5, help="Max segment duration in seconds (default: 5)")
    parser.add_argument("--device", choices=["cuda", "cpu"], default=None, help="Force device (default: auto)")
    parser.add_argument("--compute-type", choices=["float16", "int8"], default=None, help="Force compute type")
    parser.add_argument("--formats", nargs="+", default=list(DEFAULT_FORMATS), choices=ALL_FORMATS,
                        help=f"Export formats (default: {' '.join(sorted(DEFAULT_FORMATS))})")
    parser.add_argument("--initial-prompt", default=None, help="Text prompt to guide transcription")
    parser.add_argument("--high-accuracy", action="store_true", help="Higher accuracy, slower (beam_size=10)")

    args = parser.parse_args()
    args.formats = set(args.formats)

    if args.audio_file:
        if not os.path.isfile(args.audio_file):
            parser.error(f"Audio file not found: {args.audio_file}")
        transcribe_audio(args.audio_file, args)
    elif args.batch:
        if not os.path.isdir(args.batch):
            parser.error(f"Directory not found: {args.batch}")
        files = _find_audio_files(args.batch)
        run_batch(files, args)


def main():
    if len(sys.argv) <= 1:
        run_gui()
    else:
        main_cli()


if __name__ == "__main__":
    main()
